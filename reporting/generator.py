"""报告生成器 — Word / PDF / HTML 格式。"""

from html import escape
from datetime import datetime
from tempfile import NamedTemporaryFile
from typing import Optional

from models.session import Session


class ReportGenerator:
    """从分析会话或消息章节生成报告。"""

    SECTION_TITLES = [
        "分析概述",
        "专利数据总览",
        "趋势分析",
        "技术热点分析",
        "竞争对手分析",
        "技术路线图",
        "关键专利清单",
        "结论与建议",
    ]

    def __init__(self, session: Optional[Session] = None):
        self.session = session
        self.sections: list[dict] = []

    def add_section(self, title: str, content: str,
                    chart_html: str = None) -> None:
        """添加报告章节。"""
        self.sections.append({
            "title": title,
            "content": content,
            "chart_html": chart_html,
        })

    def build_from_session(self) -> None:
        """从 Session 的工具执行历史自动构建报告。"""
        if not self.session or not self.session.tool_executions:
            return

        self.add_section(
            "分析概述",
            f"会话: {self.session.name}\n"
            f"创建时间: {self.session.created_at}\n"
            f"数据集ID: {self.session.dataset_id}",
        )

        for i, exec_item in enumerate(self.session.tool_executions):
            if exec_item.status != "completed" or not exec_item.result:
                continue

            title = self.SECTION_TITLES[min(i + 1, len(self.SECTION_TITLES) - 1)]
            content = f"工具: {exec_item.tool_name}\n"
            content += f"耗时: {exec_item.duration_ms:.0f}ms\n"

            result = exec_item.result
            if hasattr(result, "data") and isinstance(result.data, list):
                content += f"数据条目: {len(result.data)}\n"
            if hasattr(result, "years"):
                content += f"年份范围: {min(result.years)}-{max(result.years)}\n"

            self.add_section(title, content, getattr(result, "chart_html", None))

    def generate_word(self, title: str = "专利分析报告",
                      sections: list[dict] = None) -> Optional[bytes]:
        """生成 Word 文档；python-docx 不可用时返回 None。"""
        try:
            from docx import Document
        except ImportError:
            return None

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        for i, sec in enumerate(sections or self.sections, 1):
            doc.add_heading(f"{i}. {sec['title']}", level=1)
            if sec.get("content"):
                doc.add_paragraph(sec["content"])
            if sec.get("chart_html"):
                doc.add_paragraph("[图表数据已生成，请在 HTML 报告中查看]")

        with NamedTemporaryFile(suffix=".docx") as file:
            doc.save(file.name)
            file.seek(0)
            return file.read()

    def generate_pdf(self, title: str = "专利分析报告",
                     sections: list[dict] = None) -> Optional[bytes]:
        """生成 PDF 报告；weasyprint 不可用时返回 None。"""
        try:
            from weasyprint import HTML
        except ImportError:
            return None

        safe_title = escape(title, quote=True)
        html_parts = [
            '<html><head><meta charset="utf-8">',
            "<style>",
            'body { font-family: "Noto Sans CJK SC", Arial, sans-serif; margin: 40px; }',
            "h1 { color: #1a1a2e; border-bottom: 2px solid #333; padding-bottom: 8px; }",
            "h2 { color: #333; }",
            "p { line-height: 1.6; }",
            ".chart { margin: 20px 0; }",
            "</style></head><body>",
            f"<h1>{safe_title}</h1>",
            f'<p>生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>',
        ]

        for i, sec in enumerate(sections or self.sections, 1):
            html_parts.append(f'<h2>{i}. {escape(str(sec["title"]), quote=True)}</h2>')
            if sec.get("content"):
                safe_content = escape(str(sec["content"]), quote=True).replace(chr(10), "<br>")
                html_parts.append(f"<p>{safe_content}</p>")
            if sec.get("chart_html"):
                html_parts.append(f'<div class="chart">{sec["chart_html"]}</div>')

        html_parts.append("</body></html>")
        return HTML(string="\n".join(html_parts)).write_pdf()

    def generate_html(self, title: str = "专利分析报告",
                      sections: list[dict] = None) -> str:
        """生成带内嵌图表的完整 HTML 报告。"""
        safe_title = escape(title, quote=True)
        html = [
            "<!DOCTYPE html>",
            '<html lang="zh-CN">',
            "<head>",
            '<meta charset="utf-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1">',
            f"<title>{safe_title}</title>",
            "<style>",
            '  body { background:#f8fafc; color:#172033; font-family: "PingFang SC","Microsoft YaHei","Noto Sans SC",Arial,sans-serif; ',
            "         padding:36px; max-width:980px; margin:0 auto; line-height:1.75; }",
            "  h1 { border-bottom:2px solid #2563eb; padding-bottom:12px; }",
            "  h2 { color:#1d4ed8; margin-top:30px; }",
            "  p { background:#fff; border:1px solid #e2e8f0; border-radius:12px; padding:18px; }",
            "  .chart { margin:20px 0; background:#fff; border-radius:12px; padding:16px; }",
            "  table { width:100%; border-collapse:collapse; font-size:12px; }",
            "  th,td { border:1px solid #e2e8f0; padding:7px; text-align:left; vertical-align:top; }",
            "  th { background:#eff6ff; color:#1e40af; }",
            "  .meta { color:#64748b; font-size:14px; background:transparent; border:0; padding:0; }",
            "</style>",
            "</head>",
            "<body>",
            f"<h1>{safe_title}</h1>",
            f'<p class="meta">生成: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>',
        ]

        for i, sec in enumerate(sections or self.sections, 1):
            html.append(f'<h2>{i}. {escape(str(sec["title"]), quote=True)}</h2>')
            if sec.get("content"):
                safe_content = escape(str(sec["content"]), quote=True).replace(chr(10), "<br>")
                html.append(f"<p>{safe_content}</p>")
            if sec.get("chart_html"):
                html.append(f'<div class="chart">{sec["chart_html"]}</div>')

        html.append("</body></html>")
        return "\n".join(html)
