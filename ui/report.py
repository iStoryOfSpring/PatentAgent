"""报告生成器 — Word / PDF 格式"""

import os
from datetime import datetime
from typing import Optional

from models.session import Session


class ReportGenerator:
    """分析报告生成。

    模板结构（8 节）:
      1. 分析概述（目标、数据范围、方法论选择依据）
      2. 专利数据总览（总量、时间跨度、主要IPC、主要申请人）
      3. 趋势分析（带图）
      4. 技术热点分析（带图）
      5. 竞争对手分析（带图）
      6. 技术路线图
      7. 关键专利清单
      8. 结论与建议
    """

    SECTION_TITLES = [
        "分析概述",
        "专利数据总览",
        "趋势分析",
        "技术热点分析",
        "竞争对手分析",
        "技术路线图",
        "关键专利清单",
        "结论与建议",
    ]

    def __init__(self, session: Optional[Session] = None):
        self.session = session
        self.sections: list[dict] = []

    def add_section(self, title: str, content: str,
                    chart_html: str = None) -> None:
        """添加报告章节"""
        self.sections.append({
            "title": title,
            "content": content,
            "chart_html": chart_html,
        })

    def build_from_session(self) -> None:
        """从 Session 的工具执行历史自动构建报告"""
        if not self.session or not self.session.tool_executions:
            return

        # 1. 分析概述
        self.add_section(
            "分析概述",
            f"会话: {self.session.name}\n"
            f"创建时间: {self.session.created_at}\n"
            f"数据集ID: {self.session.dataset_id}",
        )

        # 遍历工具执行结果构建后续章节
        for i, exec_item in enumerate(self.session.tool_executions):
            if exec_item.status != "completed" or not exec_item.result:
                continue

            title = self.SECTION_TITLES[min(i + 1, len(self.SECTION_TITLES) - 1)]
            content = f"工具: {exec_item.tool_name}\n"
            content += f"耗时: {exec_item.duration_ms:.0f}ms\n"

            result = exec_item.result
            if hasattr(result, 'data') and isinstance(result.data, list):
                content += f"数据条目: {len(result.data)}\n"
            if hasattr(result, 'years'):
                content += f"年份范围: {min(result.years)}-{max(result.years)}\n"

            chart_html = getattr(result, 'chart_html', None)
            self.add_section(title, content, chart_html)

    # ── Word 生成 ──
    def generate_word(self, title: str = "专利分析报告",
                      sections: list[dict] = None) -> Optional[bytes]:
        """生成 Word 文档。

        Returns:
            Word 文档的 bytes，或 None（python-docx 不可用时）
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            return None

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        items = sections or self.sections
        for i, sec in enumerate(items, 1):
            doc.add_heading(f"{i}. {sec['title']}", level=1)
            if sec.get('content'):
                doc.add_paragraph(sec['content'])
            if sec.get('chart_html'):
                doc.add_paragraph("[图表数据已生成，请在 HTML 报告中查看]")

        # 保存到临时文件并读取
        tmp_path = "/tmp/patent_report.docx"
        doc.save(tmp_path)
        with open(tmp_path, 'rb') as f:
            data = f.read()
        os.remove(tmp_path)
        return data

    # ── PDF 生成 ──
    def generate_pdf(self, title: str = "专利分析报告",
                     sections: list[dict] = None) -> Optional[bytes]:
        """生成 PDF 报告。

        Returns:
            PDF 文档的 bytes，或 None（weasyprint 不可用时）
        """
        try:
            from weasyprint import HTML
        except ImportError:
            return None

        # 构建 HTML
        html_parts = [
            '<html><head><meta charset="utf-8">',
            '<style>',
            'body { font-family: "Noto Sans CJK SC", Arial, sans-serif; margin: 40px; }',
            'h1 { color: #1a1a2e; border-bottom: 2px solid #333; padding-bottom: 8px; }',
            'h2 { color: #333; }',
            'p { line-height: 1.6; }',
            '.chart { margin: 20px 0; }',
            '</style></head><body>',
            f'<h1>{title}</h1>',
            f'<p>生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>',
        ]

        items = sections or self.sections
        for i, sec in enumerate(items, 1):
            html_parts.append(f'<h2>{i}. {sec["title"]}</h2>')
            if sec.get('content'):
                html_parts.append(f'<p>{sec["content"].replace(chr(10), "<br>")}</p>')
            if sec.get('chart_html'):
                html_parts.append(f'<div class="chart">{sec["chart_html"]}</div>')

        html_parts.append('</body></html>')
        html_str = '\n'.join(html_parts)

        return HTML(string=html_str).write_pdf()

    # ── HTML 报告（始终可用） ──
    def generate_html(self, title: str = "专利分析报告",
                      sections: list[dict] = None) -> str:
        """生成完整 HTML 报告（内嵌图表），含 charset 声明。"""
        items = sections or self.sections

        html = [
            '<!DOCTYPE html>',
            '<html lang="zh-CN">',
            '<head>',
            '<meta charset="utf-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1">',
            f'<title>{title}</title>',
            '<style>',
            '  body { background:#1a1a2e; color:#e0e0e0; font-family: "PingFang SC","Microsoft YaHei","Noto Sans SC",Arial,sans-serif; ',
            '         padding:30px; max-width:900px; margin:0 auto; line-height:1.8; }',
            '  h1 { border-bottom:2px solid #444; padding-bottom:10px; }',
            '  h2 { color:#FFD700; margin-top:30px; }',
            '  .chart { margin:20px 0; }',
            '  .meta { color:#888; font-size:14px; }',
            '</style>',
            '</head>',
            '<body>',
            f'<h1>{title}</h1>',
            f'<p class="meta">生成: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>',
        ]

        for i, sec in enumerate(items, 1):
            html.append(f'<h2>{i}. {sec["title"]}</h2>')
            if sec.get('content'):
                html.append(
                    f'<p>{sec["content"].replace(chr(10), "<br>")}</p>'
                )
            if sec.get('chart_html'):
                html.append(f'<div class="chart">{sec["chart_html"]}</div>')

        html.append('</body></html>')
        return '\n'.join(html)
